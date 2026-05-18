"""
文档处理器

处理 Word、PDF 等文档格式
"""

from pathlib import Path
from typing import Any, Dict, List
from .base_processor import BaseFileProcessor


class DocumentProcessor(BaseFileProcessor):
    """Word/PDF 文档处理器"""
    
    async def to_text(self, file_path: Path) -> str:
        """将文档转换为文本表示"""
        
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return await self._pdf_to_text(file_path)
        elif ext in ['.docx', '.doc']:
            return await self._word_to_text(file_path)
        else:
            raise ValueError(f"不支持的文档格式: {ext}")
    
    async def _pdf_to_text(self, file_path: Path) -> str:
        """PDF 转文本表示"""
        
        try:
            import pdfplumber
        except ImportError:
            return "[错误] 需要安装 pdfplumber: pip install pdfplumber"
        
        text_parts = [
            f"[PDF文档: {file_path.name}]\n\n"
        ]
        
        with pdfplumber.open(file_path) as pdf:
            text_parts.append(f"总页数: {len(pdf.pages)}\n\n")
            
            for i, page in enumerate(pdf.pages, 1):
                text_parts.append(f"\n{'='*50}\n")
                text_parts.append(f"第 {i} 页\n")
                text_parts.append(f"{'='*50}\n")
                
                # 提取文本
                text = page.extract_text()
                if text:
                    text_parts.append(text)
                
                # 提取表格
                tables = page.extract_tables()
                if tables:
                    for j, table in enumerate(tables, 1):
                        text_parts.append(f"\n[表格 {j}]\n")
                        text_parts.append(self._table_to_markdown(table))
        
        return "".join(text_parts)
    
    async def _word_to_text(self, file_path: Path) -> str:
        """Word 转文本表示"""
        
        try:
            from docx import Document
        except ImportError:
            return "[错误] 需要安装 python-docx: pip install python-docx"
        
        doc = Document(file_path)
        
        text_parts = [
            f"[Word文档: {file_path.name}]\n\n"
        ]
        
        # 统计信息
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        text_parts.append(f"段落数: {para_count}\n")
        text_parts.append(f"表格数: {table_count}\n\n")
        
        # 处理段落
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # 检测标题样式
            if para.style.name.startswith('Heading'):
                level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                text_parts.append(f"\n{'#' * int(level)} {para.text}\n")
            else:
                text_parts.append(f"{para.text}\n")
        
        # 处理表格
        for i, table in enumerate(doc.tables, 1):
            text_parts.append(f"\n[表格 {i}]\n")
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            text_parts.append(self._table_to_markdown(table_data))
        
        return "".join(text_parts)
    
    def _table_to_markdown(self, table: List[List]) -> str:
        """表格转 Markdown 格式"""
        if not table:
            return ""
        
        lines = []
        
        # 表头
        if len(table) > 0:
            header = table[0]
            lines.append("| " + " | ".join(str(cell) if cell else "" for cell in header) + " |")
            lines.append("| " + " | ".join("---" for _ in header) + " |")
        
        # 数据行
        for row in table[1:]:
            lines.append("| " + " | ".join(str(cell) if cell else "" for cell in row) + " |")
        
        return "\n".join(lines) + "\n"
    
    async def from_text(self, text: str, output_path: Path) -> bool:
        """从文本描述生成 Word 文档"""
        
        try:
            from docx import Document
        except ImportError:
            return False
        
        doc = Document()
        
        # 解析文本结构
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            else:
                doc.add_paragraph(line)
        
        doc.save(output_path)
        return True
    
    def get_supported_extensions(self) -> list[str]:
        return ['.pdf', '.docx', '.doc', '.rtf']
