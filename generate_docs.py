from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = r"e:\alonechat-all"

# ============================================================
# 内容定义
# ============================================================

MARKDOWN_CONTENT = r"""# Agent框架与聊天软件：对比分析及从零开发计划

> 生成日期：2026-05-10

---

## 第一章：主流Agent框架调研

### 1.1 概述

AI Agent（智能体）框架是构建基于大语言模型（LLM）的自主决策系统的中间件层。Agent = LLM + 记忆（Memory）+ 规划（Planning）+ 工具调用（Tool Use）。截至2026年，主流Agent框架已形成三足鼎立格局，同时新兴框架不断涌现。

### 1.2 核心框架介绍

#### 1.2.1 LangChain / LangGraph

| 项目 | 说明 |
|------|------|
| **开发者** | LangChain Inc. |
| **GitHub Stars** | 117k+ |
| **语言** | Python / TypeScript |
| **协议** | MIT |
| **核心理念** | 链式调用（Chain）+ 有向图（Graph）编排 |

LangChain是目前最流行的LLM应用开发框架。LangGraph是其面向Agent的扩展，将Agent执行视为一个有向图（State Graph），每个节点是一个Agent或工具调用，边是状态转换条件。

**核心特性：**
- 链式调用与管道组合
- LCEL（LangChain Expression Language）
- 丰富的工具/插件生态
- LangGraph支持状态持久化与恢复
- Human-in-the-Loop（人机协作）
- LangSmith可观测性平台

**与聊天软件的关系：**
- 提供 `ChatModel` 抽象层，统一各种LLM的聊天接口
- 可用于构建聊天机器人的RAG管道
- 通过 `CallbackHandler` 可实现实时消息流式传输
- 本身不提供通信层，需要集成到聊天应用中

#### 1.2.2 CrewAI

| 项目 | 说明 |
|------|------|
| **开发者** | CrewAI Inc. |
| **GitHub Stars** | 50k+ |
| **语言** | Python |
| **协议** | MIT |
| **核心理念** | 角色分工 + 任务驱动 + 团队协作 |

CrewAI模拟人类团队协作模式，通过角色（Role）、任务（Task）和流程（Process）驱动Agent协作。

**核心特性：**
- 高抽象度的角色定义
- Agent间任务自动分配
- 顺序/层级/协作多流程模式
- 工具集成

**与聊天软件的关系：**
- 可嵌入聊天软件作为后台编排引擎
- Agent的中间思考过程可通过聊天界面实时展示
- 适用于"聊天 + 任务执行"的复合场景

#### 1.2.3 AutoGen（Microsoft）

| 项目 | 说明 |
|------|------|
| **开发者** | Microsoft Research |
| **GitHub Stars** | 40k+ |
| **语言** | Python / .NET |
| **协议** | MIT |
| **核心理念** | 对话驱动 + 多Agent自由交流 |

AutoGen通过可定制的、可对话的Agent网络，实现Agent之间的自由交流和问题解决。其最新版本AutoGen 0.4（核心库为 `autogen-core`）进行了全面重构。

**核心特性：**
- 多Agent对话驱动
- 内置代码执行与沙箱
- 动态Agent网络
- 强类型事件驱动架构

**与聊天软件的关系：**
- 天然适合聊天场景，Agent间通过"消息"通信
- 对话驱动模式与聊天架构高度吻合
- 聊天软件可作为AutoGen的UI层

#### 1.2.4 AutoGPT

| 项目 | 说明 |
|------|------|
| **开发者** | Significant Gravitas |
| **GitHub Stars** | 150k+ |
| **语言** | Python |
| **协议** | MIT |
| **核心理念** | 自主目标分解与执行 |

AutoGPT是最早出圈的自主Agent项目，能自动将目标分解为子任务并调用工具执行。

**核心特性：**
- 自主目标分解
- 文件操作与网络访问
- 长期/短期记忆管理
- 轻量级设计

**与聊天软件的关系：**
- 适合作为聊天软件中的"自主助手"
- 可执行长期后台任务
- 通过聊天界面接收目标并汇报进度

#### 1.2.5 MetaGPT

| 项目 | 说明 |
|------|------|
| **开发者** | 深度赋智（DeepWisdom） |
| **GitHub Stars** | 60k+ |
| **语言** | Python |
| **协议** | MIT |
| **核心理念** | 结构化角色模拟 + SOP驱动 |

MetaGPT通过模拟软件公司的多角色协作（产品经理、架构师、工程师等），将SOP（标准操作流程）注入Agent协作中。

**核心特性：**
- 结构化角色定义
- 可交付物驱动的流程
- 软件开发全流程模拟

**与聊天软件的关系：**
- 可嵌入聊天软件作为团队协作助手
- 输出结构化的文档/代码到聊天界面

#### 1.2.6 Semantic Kernel（Microsoft）

| 项目 | 说明 |
|------|------|
| **开发者** | Microsoft |
| **语言** | Python / C# / Java |
| **核心理念** | 企业级AI编排 + AI服务插件化 |

Semantic Kernel是微软推出的企业级AI编排框架，深度集成Azure生态。

**核心特性：**
- Plugin系统设计成熟
- 多语言支持
- 企业级安全与合规特性
- 与Microsoft Copilot生态集成

#### 1.2.7 OpenAI Assistants API

| 项目 | 说明 |
|------|------|
| **开发者** | OpenAI |
| **协议** | 商业API |
| **核心理念** | 托管式Agent服务 |

OpenAI提供的云端Agent服务，内置代码解释器、知识检索和函数调用能力。

**核心特性：**
- 开箱即用，无需自建基础设施
- 内置RAG（知识检索）
- 托管式代码沙箱执行
- 按Token计费

---

### 1.3 框架全景对比

| 维度 | LangChain/LangGraph | CrewAI | AutoGen | AutoGPT | MetaGPT | Semantic Kernel |
|------|-------------------|--------|---------|---------|---------|----------------|
| **设计哲学** | 链/图编排 | 角色协作 | 对话驱动 | 自主执行 | SOP模拟 | 企业编排 |
| **多Agent协作** | 手动配置 | 原生支持 | 原生支持 | 不支持 | 原生支持 | 插件级支持 |
| **流程控制** | 强（状态图） | 中（流程模板） | 中（对话流） | 弱 | 强（SOP） | 强（Pipeline） |
| **工具调用** | 丰富生态 | 插件系统 | 函数注册 | 内置工具 | 内置工具 | Plugin系统 |
| **记忆管理** | 多种Memory类型 | 基础记忆 | 会话记忆 | 长短期记忆 | 结构化输出 | 语义记忆 |
| **RAG集成** | 原生支持 | 需集成 | 需集成 | 需集成 | 不内置 | 原生支持 |
| **流式输出** | 支持 | 有限 | 支持 | 有限 | 有限 | 支持 |
| **可观测性** | LangSmith | 有限 | 日志 | 日志 | 日志 | Azure Monitor |
| **学习曲线** | 陡峭 | 平缓 | 中等 | 平缓 | 中等 | 中等 |
| **生产就绪度** | 高 | 中 | 中高 | 低 | 中 | 高 |

---

### 1.4 Agent框架与聊天软件的深度融合关系

```
┌─────────────────────────────────────────────────────────┐
│                    聊天软件 (Chat App)                     │
│  ┌──────────┐  ┌──────────┐  ┌──────────────────────┐   │
│  │ 消息收发  │  │ 会话管理  │  │  Agent 集成层         │   │
│  │WebSocket │  │ 历史存储  │  │ 消息→任务·工具·规划   │   │
│  └────┬─────┘  └────┬─────┘  └──────────┬───────────┘   │
│       │             │                   │                 │
└───────┼─────────────┼───────────────────┼─────────────────┘
        │             │                   │
        ▼             ▼                   ▼
┌─────────────────────────────────────────────────────────┐
│                   Agent 框架层                            │
│  ┌──────────┐  ┌──────────┐  ┌──────────────────────┐   │
│  │ LLM 调用  │  │ 工具/插件 │  │  编排引擎             │   │
│  │ 多模型适配 │  │ 函数调用  │  │  图/链/对话驱动      │   │
│  └──────────┘  └──────────┘  └──────────────────────┘   │
│  ┌──────────┐  ┌──────────┐                              │
│  │ 记忆系统  │  │ RAG 检索  │                              │
│  │ 向量数据库 │  │ Embedding│                              │
│  └──────────┘  └──────────┘                              │
└─────────────────────────────────────────────────────────┘
```

**关键融合模式：**

1. **消息即任务：** 聊天软件中的消息天然映射为Agent的输入/任务，Agent的思考过程通过流式消息返回。
2. **会话即上下文：** 聊天会话为Agent提供长短期记忆的天然容器。
3. **群聊即多Agent协作：** 聊天群组天然支持多Agent的角色分工与协作。
4. **工具调用即消息交互：** Agent的工具调用结果可通过聊天消息格式展示。
5. **插件即对话扩展：** Agent框架的插件系统等同于聊天软件的Bot功能。

---

## 第二章：技术选型建议

### 2.1 自研 vs 采用现有框架

| 维度 | 采用现有框架 | 自研 |
|------|------------|------|
| **开发速度** | 快，周级可用 | 慢，月级起步 |
| **灵活度** | 受框架约束 | 完全可控 |
| **生态** | 丰富工具链 | 需自建 |
| **学习成本** | 需学习框架API | 无外部依赖 |
| **维护成本** | 跟随社区更新 | 全自主维护 |
| **定制化** | 有限 | 无限 |

### 2.2 推荐技术栈

#### 聊天软件

| 层次 | 技术选择 | 理由 |
|------|---------|------|
| **后端语言** | Python 3.12+ | 与Agent框架统一语言 |
| **Web框架** | FastAPI | 原生异步支持 + WebSocket |
| **实时通信** | WebSocket + Redis Pub/Sub | 高并发实时消息 |
| **数据库** | PostgreSQL + Redis | 持久化 + 缓存 |
| **消息队列** | RabbitMQ / Kafka | 异步任务处理 |
| **前端** | Next.js (React) | SSR + API路由 |
| **UI组件** | Tailwind CSS + Shadcn/ui | 快速构建界面 |

#### Agent框架（自研）

| 模块 | 技术选择 | 理由 |
|------|---------|------|
| **LLM调用层** | LiteLLM / 自研 | 统一多模型API |
| **向量数据库** | Chroma / Qdrant / Milvus | RAG记忆存储 |
| **Embedding** | 开源模型（BGE/GTE） | 私有化部署 |
| **工具执行** | 沙箱（Docker/Subprocess） | 安全执行 |
| **任务编排** | 自研DAG引擎 | 灵活可控 |
| **可观测性** | OpenTelemetry | 标准化追踪 |

---

## 第三章：从零开发计划

### 3.1 项目总览

```
Project: ChatAgent
├── chat-app/          # 聊天软件
│   ├── backend/       # Python FastAPI 后端
│   └── frontend/      # Next.js 前端
├── agent-framework/   # Agent框架（Python包）
│   ├── core/          # 核心引擎
│   ├── tools/         # 工具系统
│   ├── memory/        # 记忆系统
│   └── rag/           # RAG模块
└── integration/       # 集成层
```

### 3.2 第一阶段：聊天软件基础（第1-4周）

**目标：** 实现一个可用的实时聊天软件

| 周次 | 任务 | 交付物 |
|------|------|--------|
| Week 1 | 项目初始化 + FastAPI后端框架 | REST API基础、数据库模型 |
| Week 1 | 用户系统 | 注册/登录/JWT认证 |
| Week 2 | WebSocket实时通信 | 一对一消息、在线状态 |
| Week 2 | 消息持久化 | PostgreSQL消息存储 |
| Week 3 | 群组功能 | 群聊创建、成员管理 |
| Week 3 | Next.js前端基础 | 布局、用户界面 |
| Week 4 | 聊天UI | 消息列表、输入框、表情 |
| Week 4 | 端到端联调 | 完整聊天流程 |

**技术要点：**
- FastAPI `WebSocket` 端点管理连接生命周期
- Redis Pub/Sub 实现水平扩展的实时消息分发
- Next.js Server Actions 或 API Routes 处理HTTP请求
- React Context / Zustand 管理客户端状态

### 3.3 第二阶段：Agent框架核心（第5-10周）

**目标：** 实现一个基础的Agent框架

| 周次 | 任务 | 交付物 |
|------|------|--------|
| Week 5 | 框架架构设计 | 核心接口定义（Agent, Tool, Memory） |
| Week 5 | LLM调用层 | LiteLLM集成，多模型适配 |
| Week 6 | 工具系统 | 工具注册/发现/执行机制 |
| Week 6 | Agent核心引擎 | ReAct循环、思考-行动-观察 |
| Week 7 | 记忆系统 | 短期（会话）记忆 + 长期（向量）记忆 |
| Week 7 | RAG模块 | 文档加载、分块、Embedding、检索 |
| Week 8 | 任务编排引擎 | 顺序/并行/DAG工作流 |
| Week 8 | 多Agent协作 | Agent间消息传递、角色分配 |
| Week 9 | 可观测性 | Logging、Tracing、Token计数 |
| Week 9 | 安全性 | 工具沙箱、速率限制、内容过滤 |
| Week 10 | 单元测试 + 文档 | 核心模块测试覆盖率 > 80% |

**Agent框架核心架构：**

```
Agent框架核心接口
├── BaseAgent
│   ├── perceive(input) -> Perception
│   ├── plan(perception) -> Plan
│   ├── act(plan) -> ActionResult[]
│   └── reflect(result) -> Reflection
├── BaseTool
│   ├── name: str
│   ├── description: str
│   ├── parameters: JSONSchema
│   └── execute(params) -> Any
├── BaseMemory
│   ├── add(entry)
│   ├── query(Query) -> List[Entry]
│   └── clear()
├── BaseLLM
│   ├── chat(messages) -> Message
│   └── chat_stream(messages) -> Stream[Message]
└── Orchestrator
    ├── run(agent, task) -> Result
    ├── run_multi(agents, task) -> Result
    └── run_workflow(graph) -> Result
```

### 3.4 第三阶段：聊天 + Agent 融合（第11-14周）

**目标：** 将Agent能力嵌入聊天软件

| 周次 | 任务 | 交付物 |
|------|------|--------|
| Week 11 | Agent集成层 | 聊天消息→Agent任务转换 |
| Week 11 | 流式响应 | Agent思考过程实时展示 |
| Week 12 | 会话Agent | 每个聊天会话绑定Agent上下文 |
| Week 12 | 群聊Agent | 群聊中多Agent协作 |
| Week 13 | Agent商店 | 插件/工具/Agent模板市场 |
| Week 13 | 管理后台 | Agent监控、日志、配置 |
| Week 14 | 压力测试 + 优化 | 性能报告、优化方案 |

### 3.5 第四阶段：生产化（第15-18周）

| 周次 | 任务 |
|------|------|
| Week 15 | Docker容器化 + Docker Compose编排 |
| Week 15 | CI/CD流水线 |
| Week 16 | 水平扩展（WebSocket集群 + Agent工作节点） |
| Week 16 | 数据库连接池 + 缓存优化 |
| Week 17 | 安全审计 + 渗透测试 |
| Week 17 | 多语言国际化 |
| Week 18 | 公测 + 反馈迭代 |

---

## 第四章：架构设计要点

### 4.1 聊天软件架构

```
┌────────────────────────────────────────────────────┐
│                  客户端 (Next.js)                    │
│  ┌─────────────┐  ┌─────────────┐                  │
│  │   聊天UI     │  │  Agent UI   │                  │
│  │ 消息列表/输入 │  │ 思考过程/工具│                  │
│  └──────┬──────┘  └──────┬──────┘                  │
│         │                │                          │
│    ┌────┴────────────────┴────┐                     │
│    │     WebSocket Client     │                     │
│    └────────────┬─────────────┘                     │
└─────────────────┼───────────────────────────────────┘
                  │ WSS
┌─────────────────┼───────────────────────────────────┐
│          反向代理 / 负载均衡 (Nginx)                  │
└─────────────────┼───────────────────────────────────┘
                  │
┌─────────────────┼───────────────────────────────────┐
│          FastAPI 应用服务器                            │
│  ┌──────────────┴──────────────┐                     │
│  │     WebSocket Manager       │                     │
│  │  连接管理 / 房间路由 / 心跳  │                     │
│  └──────────────┬──────────────┘                     │
│  ┌──────────────┴──────────────┐                     │
│  │      消息处理管道            │                     │
│  │ 接收 → 验证 → 存储 → 广播   │                     │
│  └──────────────┬──────────────┘                     │
│  ┌──────────────┴──────────────┐                     │
│  │      Agent 集成层           │                     │
│  │  消息→Agent任务 / 流式返回  │                     │
│  └──────────────┬──────────────┘                     │
│  ┌──────────────┴──────────────┐                     │
│  │      REST API              │                     │
│  │  用户/群组/文件/历史记录    │                     │
│  └─────────────────────────────┘                     │
└─────────────────┬───────────────────────────────────┘
                  │
┌─────────────────┼───────────────────────────────────┐
│      数据层                                          │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐          │
│  │PostgreSQL│  │  Redis   │  │ 对象存储  │          │
│  │消息/用户  │  │ 会话/缓存 │  │ 文件/图片 │          │
│  └──────────┘  └──────────┘  └──────────┘          │
└─────────────────────────────────────────────────────┘
```

### 4.2 Agent框架架构

```
┌────────────────────────────────────────────────────┐
│                    Agent 框架                         │
│                                                      │
│  ┌──────────────────────────────────────────────┐   │
│  │              编排引擎 (Orchestrator)          │   │
│  │  顺序执行 │ DAG工作流 │ 多Agent协商          │   │
│  └────┬─────┬──────────┬──────────┬─────────────┘   │
│       │     │          │          │                   │
│  ┌────┴┐ ┌──┴──┐ ┌────┴────┐ ┌───┴────┐             │
│  │Agent│ │Tool │ │ Memory  │ │  LLM   │             │
│  │执行器│ │系统  │ │ 记忆系统 │ │ 调用层  │             │
│  └─────┘ └─────┘ └─────────┘ └────────┘             │
│  ┌──────────────────────────────────────────────┐   │
│  │               RAG 管道                         │   │
│  │  文档解析 → 分块 → Embedding → 检索 → 重排序 │   │
│  └──────────────────────────────────────────────┘   │
│  ┌──────────────────────────────────────────────┐   │
│  │           可观测性 & 安全                       │   │
│  │  OpenTelemetry │ 沙箱执行 │ 速率限制 │ 审计日志 │   │
│  └──────────────────────────────────────────────┘   │
└────────────────────────────────────────────────────┘
```

### 4.3 聊天与Agent融合的关键设计

1. **消息协议扩展：** 在标准聊天消息格式中增加 `agent_meta` 字段，用于携带Agent的思考过程、工具调用状态等信息。
2. **异步响应模式：** Agent任务可能耗时较长，采用"任务提交 → 消息通知 → 结果推送"的异步模式。
3. **上下文管理：** 聊天会话的完整历史作为Agent的短期记忆，超过窗口大小后压缩或向量化存储为长期记忆。
4. **权限控制：** Agent可调用的工具受聊天室/用户权限约束。

---

## 第五章：风险与应对

| 风险 | 影响 | 应对方案 |
|------|------|----------|
| LLM API成本过高 | 运营成本失控 | 采用多模型路由（简单任务用小模型），缓存复用 |
| Agent执行不可控 | 用户体验差 | 加入人工确认节点，限制工具权限 |
| 实时消息延迟 | 聊天体验差 | WebSocket连接池优化，Redis集群 |
| 数据库性能瓶颈 | 消息获取变慢 | 读写分离，消息分表，CDN缓存 |
| 安全问题（Prompt注入） | 系统被攻击 | 输入过滤，工具沙箱，权限最小化 |

---

## 第六章：总结与路线图

### 6.1 核心结论

1. **Agent框架与聊天软件天然互补：** 聊天软件提供实时通信层，Agent框架提供智能决策层，两者融合可实现"智能对话即服务"。
2. **LangChain生态最成熟，CrewAI多Agent协作最便捷，AutoGen对话驱动与聊天场景最契合。**
3. **自研Agent框架的建议：** 初期可参考LangChain的设计哲学，但采用更轻量、更模块化的架构，避免过度抽象。
4. **从零开发建议分4个阶段（共18周）：** 聊天基础 → Agent核心 → 融合集成 → 生产化。

### 6.2 技术栈汇总

| 角色 | 推荐技术 |
|------|---------|
| 后端语言 | Python 3.12+ |
| Web框架 | FastAPI |
| 实时通信 | WebSocket + Redis Pub/Sub |
| 数据库 | PostgreSQL (主) + Redis (缓存) + 向量数据库 (RAG) |
| 前端框架 | Next.js 14+ (React) |
| UI框架 | Tailwind CSS + Shadcn/ui |
| LLM调用 | LiteLLM (多模型适配) |
| 容器化 | Docker + Docker Compose |
| 可观测性 | OpenTelemetry |

### 6.3 项目路线图

```
Week 1-4:   聊天软件基础功能
    └── 实时消息、用户系统、群组
Week 5-10:  Agent框架核心
    └── LLM调用、工具系统、记忆、RAG、多Agent编排
Week 11-14: 聊天+Agent融合
    └── 流式响应、会话Agent、群聊Agent、管理后台
Week 15-18: 生产化
    └── 容器化、CI/CD、水平扩展、安全审计
```
"""


def create_docx():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ---- 封面标题 ----
    title = doc.add_heading('Agent框架与聊天软件', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('对比分析及从零开发计划')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('生成日期：2026-05-10\n技术栈：Python 3.12+ / FastAPI / Next.js / React').font.size = Pt(10)

    doc.add_page_break()

    # ---- 目录占位 ----
    doc.add_heading('目录', level=1)
    toc_items = [
        '第一章 主流Agent框架调研',
        '  1.1 概述',
        '  1.2 核心框架介绍（LangChain、CrewAI、AutoGen、AutoGPT、MetaGPT、Semantic Kernel、OpenAI Assistants API）',
        '  1.3 框架全景对比',
        '  1.4 Agent框架与聊天软件的深度融合关系',
        '第二章 技术选型建议',
        '  2.1 自研 vs 采用现有框架',
        '  2.2 推荐技术栈',
        '第三章 从零开发计划',
        '  3.1 项目总览',
        '  3.2 第一阶段：聊天软件基础（第1-4周）',
        '  3.3 第二阶段：Agent框架核心（第5-10周）',
        '  3.4 第三阶段：聊天 + Agent融合（第11-14周）',
        '  3.5 第四阶段：生产化（第15-18周）',
        '第四章 架构设计要点',
        '  4.1 聊天软件架构',
        '  4.2 Agent框架架构',
        '  4.3 聊天与Agent融合的关键设计',
        '第五章 风险与应对',
        '第六章 总结与路线图',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.space_after = Pt(2)

    doc.add_page_break()

    # ========== 第一章 ==========
    doc.add_heading('第一章 主流Agent框架调研', level=1)

    doc.add_heading('1.1 概述', level=2)
    doc.add_paragraph(
        'AI Agent（智能体）框架是构建基于大语言模型（LLM）的自主决策系统的中间件层。'
        '其核心公式为：Agent = LLM + 记忆（Memory） + 规划（Planning） + 工具调用（Tool Use）。'
        '截至2026年，主流Agent框架已形成三足鼎立格局，同时新兴框架不断涌现。'
    )

    # ---- 1.2 核心框架 ----
    doc.add_heading('1.2 核心框架介绍', level=2)

    frameworks = [
        ('1.2.1 LangChain / LangGraph', [
            '开发者：LangChain Inc. | GitHub Stars: 117k+ | 语言：Python / TypeScript | 协议：MIT',
            '核心理念：链式调用（Chain）+ 有向图（Graph）编排',
            '',
            'LangChain是目前最流行的LLM应用开发框架。LangGraph是其面向Agent的扩展，将Agent执行视为一个有向图（State Graph），每个节点是一个Agent或工具调用，边是状态转换条件。',
            '',
            '核心特性：',
            '  - 链式调用与管道组合，LCEL（LangChain Expression Language）',
            '  - 丰富的工具/插件生态',
            '  - LangGraph支持状态持久化与恢复',
            '  - Human-in-the-Loop（人机协作）',
            '  - LangSmith可观测性平台',
            '',
            '与聊天软件的关系：',
            '  - 提供ChatModel抽象层，统一各种LLM的聊天接口',
            '  - 可用于构建聊天机器人的RAG管道',
            '  - 通过CallbackHandler可实现实时消息流式传输',
            '  - 本身不提供通信层，需要集成到聊天应用中',
        ]),
        ('1.2.2 CrewAI', [
            '开发者：CrewAI Inc. | GitHub Stars: 50k+ | 语言：Python | 协议：MIT',
            '核心理念：角色分工 + 任务驱动 + 团队协作',
            '',
            'CrewAI模拟人类团队协作模式，通过角色（Role）、任务（Task）和流程（Process）驱动Agent协作。',
            '',
            '核心特性：',
            '  - 高抽象度的角色定义',
            '  - Agent间任务自动分配',
            '  - 顺序/层级/协作多流程模式',
            '',
            '与聊天软件的关系：',
            '  - 可嵌入聊天软件作为后台编排引擎',
            '  - Agent的中间思考过程可通过聊天界面实时展示',
            '  - 适用于"聊天 + 任务执行"的复合场景',
        ]),
        ('1.2.3 AutoGen（Microsoft）', [
            '开发者：Microsoft Research | GitHub Stars: 40k+ | 语言：Python / .NET | 协议：MIT',
            '核心理念：对话驱动 + 多Agent自由交流',
            '',
            'AutoGen通过可定制的、可对话的Agent网络，实现Agent之间的自由交流和问题解决。',
            '',
            '核心特性：',
            '  - 多Agent对话驱动，内置代码执行与沙箱',
            '  - 动态Agent网络',
            '  - 强类型事件驱动架构',
            '',
            '与聊天软件的关系：',
            '  - 天然适合聊天场景，Agent间通过"消息"通信',
            '  - 对话驱动模式与聊天架构高度吻合',
            '  - 聊天软件可作为AutoGen的UI层',
        ]),
        ('1.2.4 AutoGPT', [
            '开发者：Significant Gravitas | GitHub Stars: 150k+ | 语言：Python | 协议：MIT',
            '核心理念：自主目标分解与执行',
            '',
            'AutoGPT是最早出圈的自主Agent项目，能自动将目标分解为子任务并调用工具执行。',
            '',
            '核心特性：',
            '  - 自主目标分解，文件操作与网络访问',
            '  - 长期/短期记忆管理',
            '  - 轻量级设计',
            '',
            '与聊天软件的关系：',
            '  - 适合作为聊天软件中的"自主助手"',
            '  - 可执行长期后台任务',
            '  - 通过聊天界面接收目标并汇报进度',
        ]),
        ('1.2.5 MetaGPT', [
            '开发者：深度赋智（DeepWisdom）| GitHub Stars: 60k+ | 语言：Python | 协议：MIT',
            '核心理念：结构化角色模拟 + SOP驱动',
            '',
            'MetaGPT通过模拟软件公司的多角色协作，将SOP（标准操作流程）注入Agent协作中。',
            '',
            '核心特性：',
            '  - 结构化角色定义',
            '  - 可交付物驱动的流程',
            '  - 软件开发全流程模拟',
        ]),
        ('1.2.6 Semantic Kernel（Microsoft）', [
            '开发者：Microsoft | 语言：Python / C# / Java',
            '核心理念：企业级AI编排 + AI服务插件化',
            '',
            '核心特性：',
            '  - Plugin系统设计成熟，多语言支持',
            '  - 企业级安全与合规特性',
            '  - 与Microsoft Copilot生态集成',
        ]),
        ('1.2.7 OpenAI Assistants API', [
            '开发者：OpenAI | 协议：商业API',
            '核心理念：托管式Agent服务',
            '',
            '核心特性：',
            '  - 开箱即用，无需自建基础设施',
            '  - 内置RAG（知识检索），托管式代码沙箱执行',
            '  - 按Token计费',
        ]),
    ]

    for title_text, lines in frameworks:
        doc.add_heading(title_text, level=3)
        for line in lines:
            if line == '':
                continue
            p = doc.add_paragraph(line)
            p.space_after = Pt(1)
            p.space_before = Pt(1)

    # ---- 1.3 全景对比表格 ----
    doc.add_heading('1.3 框架全景对比', level=2)

    table = doc.add_table(rows=11, cols=7, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['维度', 'LangChain/Graph', 'CrewAI', 'AutoGen', 'AutoGPT', 'MetaGPT', 'Semantic Kernel']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    data = [
        ['设计哲学', '链/图编排', '角色协作', '对话驱动', '自主执行', 'SOP模拟', '企业编排'],
        ['多Agent协作', '手动配置', '原生支持', '原生支持', '不支持', '原生支持', '插件级'],
        ['流程控制', '强（状态图）', '中（流程模板）', '中（对话流）', '弱', '强（SOP）', '强（Pipeline）'],
        ['工具调用', '丰富生态', '插件系统', '函数注册', '内置工具', '内置工具', 'Plugin系统'],
        ['记忆管理', '多种Memory', '基础记忆', '会话记忆', '长短期记忆', '结构化输出', '语义记忆'],
        ['RAG集成', '原生支持', '需集成', '需集成', '需集成', '不内置', '原生支持'],
        ['流式输出', '支持', '有限', '支持', '有限', '有限', '支持'],
        ['可观测性', 'LangSmith', '有限', '日志', '日志', '日志', 'Azure Monitor'],
        ['学习曲线', '陡峭', '平缓', '中等', '平缓', '中等', '中等'],
        ['生产就绪度', '高', '中', '中高', '低', '中', '高'],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    # ---- 1.4 融合关系 ----
    doc.add_heading('1.4 Agent框架与聊天软件的深度融合关系', level=2)
    doc.add_paragraph(
        'Agent框架与聊天软件之间存在着天然的互补关系。聊天软件提供实时通信层和用户界面，'
        'Agent框架提供智能决策和任务执行能力。两者融合可实现"智能对话即服务"的愿景。'
    )

    doc.add_heading('关键融合模式', level=3)
    patterns = [
        ('消息即任务：', '聊天软件中的消息天然映射为Agent的输入/任务，Agent的思考过程通过流式消息返回。'),
        ('会话即上下文：', '聊天会话为Agent提供长短期记忆的天然容器，完整对话历史是Agent上下文的理想来源。'),
        ('群聊即多Agent协作：', '聊天群组天然支持多Agent的角色分工与协作，每个Agent可扮演群聊中的特定角色。'),
        ('工具调用即消息交互：', 'Agent的工具调用结果可通过聊天消息格式展示，如卡片、图表、代码块等。'),
        ('插件即对话扩展：', 'Agent框架的插件系统等同于聊天软件的Bot功能，两者可统一为"技能"市场。'),
    ]
    for title_text, desc in patterns:
        p = doc.add_paragraph()
        run_bold = p.add_run(title_text)
        run_bold.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========== 第二章 ==========
    doc.add_heading('第二章 技术选型建议', level=1)

    doc.add_heading('2.1 自研 vs 采用现有框架', level=2)

    table2 = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    h2 = ['维度', '采用现有框架', '自研']
    for i, h in enumerate(h2):
        table2.rows[0].cells[i].text = h
        for paragraph in table2.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    data2 = [
        ['开发速度', '快，周级可用', '慢，月级起步'],
        ['灵活度', '受框架约束', '完全可控'],
        ['生态', '丰富工具链', '需自建'],
        ['学习成本', '需学习框架API', '无外部依赖'],
        ['维护成本', '跟随社区更新', '全自主维护'],
        ['定制化', '有限', '无限'],
    ]
    for r_idx, row in enumerate(data2):
        for c_idx, val in enumerate(row):
            table2.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_heading('2.2 推荐技术栈', level=2)

    doc.add_heading('聊天软件', level=3)
    table3 = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    h3 = ['层次', '技术选择', '理由']
    for i, h in enumerate(h3):
        table3.rows[0].cells[i].text = h
        for paragraph in table3.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    chat_stack = [
        ['后端语言', 'Python 3.12+', '与Agent框架统一语言'],
        ['Web框架', 'FastAPI', '原生异步支持 + WebSocket'],
        ['实时通信', 'WebSocket + Redis Pub/Sub', '高并发实时消息'],
        ['数据库', 'PostgreSQL + Redis', '持久化 + 缓存'],
        ['消息队列', 'RabbitMQ / Kafka', '异步任务处理'],
        ['前端', 'Next.js (React)', 'SSR + API路由'],
    ]
    for r_idx, row in enumerate(chat_stack):
        for c_idx, val in enumerate(row):
            table3.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_heading('Agent框架（自研）', level=3)
    table4 = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    h4 = ['模块', '技术选择', '理由']
    for i, h in enumerate(h4):
        table4.rows[0].cells[i].text = h
        for paragraph in table4.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    agent_stack = [
        ['LLM调用层', 'LiteLLM / 自研', '统一多模型API'],
        ['向量数据库', 'Chroma / Qdrant / Milvus', 'RAG记忆存储'],
        ['Embedding', '开源模型（BGE/GTE）', '私有化部署'],
        ['工具执行', '沙箱（Docker/Subprocess）', '安全执行'],
        ['任务编排', '自研DAG引擎', '灵活可控'],
        ['可观测性', 'OpenTelemetry', '标准化追踪'],
    ]
    for r_idx, row in enumerate(agent_stack):
        for c_idx, val in enumerate(row):
            table4.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_page_break()

    # ========== 第三章 ==========
    doc.add_heading('第三章 从零开发计划', level=1)

    doc.add_heading('3.1 项目总览', level=2)
    doc.add_paragraph(
        '项目名称 ChatAgent，包含三个子项目：chat-app（聊天软件）、agent-framework（Agent框架Python包）、'
        'integration（集成层）。'
    )
    p = doc.add_paragraph()
    p.add_run('ChatAgent 项目结构：').bold = True
    structure = [
        'chat-app/          # 聊天软件',
        '  ├── backend/     # Python FastAPI 后端',
        '  └── frontend/    # Next.js 前端',
        'agent-framework/   # Agent框架（Python包）',
        '  ├── core/        # 核心引擎',
        '  ├── tools/       # 工具系统',
        '  ├── memory/      # 记忆系统',
        '  └── rag/         # RAG模块',
        'integration/       # 集成层',
    ]
    for line in structure:
        doc.add_paragraph(line, style='List')

    # ---- 3.2 ----
    doc.add_heading('3.2 第一阶段：聊天软件基础（第1-4周）', level=2)
    doc.add_paragraph('目标：实现一个可用的实时聊天软件')

    table5 = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['周次', '任务', '交付物']):
        table5.rows[0].cells[i].text = h
        for paragraph in table5.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    phase1 = [
        ['Week 1', '项目初始化 + FastAPI后端框架', 'REST API基础、数据库模型'],
        ['Week 1', '用户系统', '注册/登录/JWT认证'],
        ['Week 2', 'WebSocket实时通信', '一对一消息、在线状态'],
        ['Week 2', '消息持久化', 'PostgreSQL消息存储'],
        ['Week 3', '群组功能', '群聊创建、成员管理'],
        ['Week 3', 'Next.js前端基础', '布局、用户界面'],
        ['Week 4', '聊天UI', '消息列表、输入框、表情'],
        ['Week 4', '端到端联调', '完整聊天流程'],
    ]
    for r_idx, row in enumerate(phase1):
        for c_idx, val in enumerate(row):
            table5.rows[r_idx + 1].cells[c_idx].text = val

    p = doc.add_paragraph()
    p.add_run('技术要点：').bold = True
    doc.add_paragraph('FastAPI WebSocket端点管理连接生命周期', style='List Bullet')
    doc.add_paragraph('Redis Pub/Sub实现水平扩展的实时消息分发', style='List Bullet')
    doc.add_paragraph('Next.js Server Actions或API Routes处理HTTP请求', style='List Bullet')
    doc.add_paragraph('React Context / Zustand管理客户端状态', style='List Bullet')

    # ---- 3.3 ----
    doc.add_heading('3.3 第二阶段：Agent框架核心（第5-10周）', level=2)
    doc.add_paragraph('目标：实现一个基础的Agent框架')

    table6 = doc.add_table(rows=11, cols=3, style='Light Grid Accent 1')
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['周次', '任务', '交付物']):
        table6.rows[0].cells[i].text = h
        for paragraph in table6.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    phase2 = [
        ['Week 5', '框架架构设计', '核心接口定义（Agent, Tool, Memory）'],
        ['Week 5', 'LLM调用层', 'LiteLLM集成，多模型适配'],
        ['Week 6', '工具系统', '工具注册/发现/执行机制'],
        ['Week 6', 'Agent核心引擎', 'ReAct循环、思考-行动-观察'],
        ['Week 7', '记忆系统', '短期（会话）+ 长期（向量）记忆'],
        ['Week 7', 'RAG模块', '文档加载、分块、Embedding、检索'],
        ['Week 8', '任务编排引擎', '顺序/并行/DAG工作流'],
        ['Week 8', '多Agent协作', 'Agent间消息传递、角色分配'],
        ['Week 9', '可观测性 + 安全', 'Logging、Tracing、工具沙箱'],
        ['Week 10', '单元测试 + 文档', '核心模块测试覆盖率 > 80%'],
    ]
    for r_idx, row in enumerate(phase2):
        for c_idx, val in enumerate(row):
            table6.rows[r_idx + 1].cells[c_idx].text = val

    p = doc.add_paragraph()
    p.add_run('Agent框架核心接口定义：').bold = True
    interfaces = [
        'BaseAgent: perceive() → plan() → act() → reflect()',
        'BaseTool: name + description + parameters + execute()',
        'BaseMemory: add() + query() + clear()',
        'BaseLLM: chat() + chat_stream()',
        'Orchestrator: run() / run_multi() / run_workflow()',
    ]
    for line in interfaces:
        doc.add_paragraph(line, style='List Bullet')

    # ---- 3.4 ----
    doc.add_heading('3.4 第三阶段：聊天 + Agent融合（第11-14周）', level=2)
    table7 = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
    table7.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['周次', '任务', '交付物']):
        table7.rows[0].cells[i].text = h
        for paragraph in table7.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    phase3 = [
        ['Week 11', 'Agent集成层', '聊天消息→Agent任务转换'],
        ['Week 11', '流式响应', 'Agent思考过程实时展示'],
        ['Week 12', '会话Agent', '每个聊天会话绑定Agent上下文'],
        ['Week 12', '群聊Agent', '群聊中多Agent协作'],
        ['Week 13', 'Agent商店 + 管理后台', '插件/工具/Agent模板市场'],
        ['Week 14', '压力测试 + 优化', '性能报告、优化方案'],
    ]
    for r_idx, row in enumerate(phase3):
        for c_idx, val in enumerate(row):
            table7.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 3.5 ----
    doc.add_heading('3.5 第四阶段：生产化（第15-18周）', level=2)
    table8 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
    table8.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['周次', '任务']):
        table8.rows[0].cells[i].text = h
        for paragraph in table8.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    phase4 = [
        ['Week 15', 'Docker容器化 + Docker Compose编排 + CI/CD流水线'],
        ['Week 16', '水平扩展（WebSocket集群 + Agent工作节点）+ 缓存优化'],
        ['Week 17', '安全审计 + 渗透测试 + 多语言国际化'],
        ['Week 18', '公测 + 反馈迭代'],
    ]
    for r_idx, row in enumerate(phase4):
        for c_idx, val in enumerate(row):
            table8.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_page_break()

    # ========== 第四章 ==========
    doc.add_heading('第四章 架构设计要点', level=1)

    doc.add_heading('4.1 聊天软件架构', level=2)
    arch_chat = [
        '客户端层 (Next.js): 聊天UI + Agent UI → WebSocket Client',
        '反向代理层: Nginx 负载均衡',
        '应用层 (FastAPI):',
        '  - WebSocket Manager: 连接管理 / 房间路由 / 心跳',
        '  - 消息处理管道: 接收 → 验证 → 存储 → 广播',
        '  - Agent集成层: 消息→Agent任务 / 流式返回',
        '  - REST API: 用户/群组/文件/历史记录',
        '数据层: PostgreSQL (消息/用户) + Redis (会话/缓存) + 对象存储 (文件/图片)',
    ]
    for line in arch_chat:
        doc.add_paragraph(line)

    doc.add_heading('4.2 Agent框架架构', level=2)
    arch_agent = [
        '编排引擎 (Orchestrator): 顺序执行 / DAG工作流 / 多Agent协商',
        'Agent执行器: ReAct循环',
        '工具系统: 工具注册/发现/沙箱执行',
        '记忆系统: 短期记忆 + 长期记忆（向量化）',
        'LLM调用层: 多模型适配（LiteLLM）',
        'RAG管道: 文档解析 → 分块 → Embedding → 检索 → 重排序',
        '可观测性 & 安全: OpenTelemetry / 沙箱执行 / 速率限制 / 审计日志',
    ]
    for line in arch_agent:
        doc.add_paragraph(line)

    doc.add_heading('4.3 聊天与Agent融合的关键设计', level=2)
    p = doc.add_paragraph()
    p.add_run('1. 消息协议扩展：').bold = True
    p.add_run('在标准聊天消息格式中增加agent_meta字段，用于携带Agent的思考过程、工具调用状态等信息。')
    p = doc.add_paragraph()
    p.add_run('2. 异步响应模式：').bold = True
    p.add_run('Agent任务可能耗时较长，采用"任务提交 → 消息通知 → 结果推送"的异步模式。')
    p = doc.add_paragraph()
    p.add_run('3. 上下文管理：').bold = True
    p.add_run('聊天会话的完整历史作为Agent的短期记忆，超过窗口大小后压缩或向量化存储为长期记忆。')
    p = doc.add_paragraph()
    p.add_run('4. 权限控制：').bold = True
    p.add_run('Agent可调用的工具受聊天室/用户权限约束。')

    doc.add_page_break()

    # ========== 第五章 ==========
    doc.add_heading('第五章 风险与应对', level=1)

    table9 = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
    table9.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['风险', '影响', '应对方案']):
        table9.rows[0].cells[i].text = h
        for paragraph in table9.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    risks = [
        ['LLM API成本过高', '运营成本失控', '多模型路由（简单任务用小模型），缓存复用'],
        ['Agent执行不可控', '用户体验差', '加入人工确认节点，限制工具权限'],
        ['实时消息延迟', '聊天体验差', 'WebSocket连接池优化，Redis集群'],
        ['数据库性能瓶颈', '消息获取变慢', '读写分离，消息分表，CDN缓存'],
        ['安全问题（Prompt注入）', '系统被攻击', '输入过滤，工具沙箱，权限最小化'],
    ]
    for r_idx, row in enumerate(risks):
        for c_idx, val in enumerate(row):
            table9.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_page_break()

    # ========== 第六章 ==========
    doc.add_heading('第六章 总结与路线图', level=1)

    doc.add_heading('6.1 核心结论', level=2)
    conclusions = [
        'Agent框架与聊天软件天然互补：聊天软件提供实时通信层，Agent框架提供智能决策层，两者融合可实现"智能对话即服务"。',
        'LangChain生态最成熟，CrewAI多Agent协作最便捷，AutoGen对话驱动与聊天场景最契合。',
        '自研Agent框架初期可参考LangChain的设计哲学，但采用更轻量、更模块化的架构，避免过度抽象。',
        '从零开发建议分4个阶段（共18周）：聊天基础 → Agent核心 → 融合集成 → 生产化。',
    ]
    for c in conclusions:
        doc.add_paragraph(c, style='List Number')

    doc.add_heading('6.2 技术栈汇总', level=2)
    table10 = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
    table10.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['角色', '推荐技术']):
        table10.rows[0].cells[i].text = h
        for paragraph in table10.rows[0].cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    stack_summary = [
        ['后端语言', 'Python 3.12+'],
        ['Web框架', 'FastAPI'],
        ['实时通信', 'WebSocket + Redis Pub/Sub'],
        ['数据库', 'PostgreSQL (主) + Redis (缓存) + 向量数据库 (RAG)'],
        ['前端框架', 'Next.js 14+ (React)'],
        ['LLM调用', 'LiteLLM (多模型适配)'],
        ['容器化', 'Docker + Docker Compose'],
    ]
    for r_idx, row in enumerate(stack_summary):
        for c_idx, val in enumerate(row):
            table10.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_heading('6.3 项目路线图', level=2)
    roadmap = [
        'Week 1-4:    聊天软件基础功能 — 实时消息、用户系统、群组',
        'Week 5-10:   Agent框架核心 — LLM调用、工具系统、记忆、RAG、多Agent编排',
        'Week 11-14:  聊天+Agent融合 — 流式响应、会话Agent、群聊Agent、管理后台',
        'Week 15-18:  生产化 — 容器化、CI/CD、水平扩展、安全审计',
    ]
    for line in roadmap:
        doc.add_paragraph(line)

    # 保存
    docx_path = os.path.join(OUTPUT_DIR, 'Agent框架与聊天软件_对比分析及开发计划.docx')
    doc.save(docx_path)
    print(f'DOCX saved: {docx_path}')
    return docx_path


def save_markdown():
    md_path = os.path.join(OUTPUT_DIR, 'Agent框架与聊天软件_对比分析及开发计划.md')
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(MARKDOWN_CONTENT)
    print(f'MD saved: {md_path}')
    return md_path


if __name__ == '__main__':
    create_docx()
    save_markdown()
